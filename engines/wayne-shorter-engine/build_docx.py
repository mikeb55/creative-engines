"""Build User_Guide.docx from User_Guide.md. Requires: pip install python-docx."""

import os

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    raise SystemExit(1)

HERE = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(HERE, "docs", "User_Guide.md")
DOCX_PATH = os.path.join(HERE, "docs", "User_Guide.docx")


def main():
    with open(MD_PATH, "r", encoding="utf-8") as f:
        lines = f.readlines()
    doc = Document()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("- "):
            p = doc.add_paragraph()
            p.add_run(stripped[2:])
        elif stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                doc.add_paragraph("\n".join(code_lines).rstrip(), style="Normal")
            if i < len(lines):
                i += 1
            i -= 1
        else:
            doc.add_paragraph(stripped)
        i += 1
    doc.save(DOCX_PATH)
    print(f"Saved {DOCX_PATH}")


if __name__ == "__main__":
    main()
